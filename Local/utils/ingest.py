import hashlib, time, os
import uuid
from typing import List, Dict, Any
from pypdf import PdfReader
import camelot
import pdfplumber
from docx import Document as Docx
from markdown_it import MarkdownIt
from sentence_transformers import SentenceTransformer
from qdrant_client import QdrantClient
from qdrant_client.http.models import VectorParams, Distance, HnswConfig, PointStruct,HnswConfigDiff
from qdrant_client.http import models
from whoosh import index as windex
from whoosh.fields import Schema, TEXT, ID
from utils.config import RAGConfig

Cfg = RAGConfig()

def sha1(s: str) -> str:
    return hashlib.sha1(s.encode("utf-8", errors="ignore")).hexdigest()

def now() -> int:
    return int(time.time())

def read_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        reader = PdfReader(path)
        pages = []
        for i, p in enumerate(reader.pages, start=1):
            try:
                pages.append((i, p.extract_text() or ""))
            except Exception:
                pages.append((i, ""))
        return "\\n".join([f"\\n[[PAGE:{i}]]\\n{t}" for i, t in pages])
    elif ext in {".md", ".markdown"}:
        return open(path, "r", encoding="utf-8", errors="ignore").read()
    elif ext in {".txt"}:
        return open(path, "r", encoding="utf-8", errors="ignore").read()
    elif ext in {".docx"}:
        d = Docx(path)
        return "\\n".join(p.text for p in d.paragraphs)
    else:
        raise ValueError(f"Unsupported format: {ext}")
    
def read_blocks_with_tables(path: str) -> List[Dict[str, Any]]:
    """
    Return a unified list of blocks with tables separated:
    - {"type":"text","page":int,"heading":str,"text":str}
    - {"type":"table","page":int,"heading":str,"rows":[...],"n_cols":int,"table_id":str}
    """
    ext = os.path.splitext(path)[1].lower()
    blocks: List[Dict[str, Any]] = []

    if ext == ".pdf":
        # Text (existing approach)
        reader = PdfReader(path)
        pages = []
        for i, p in enumerate(reader.pages, start=1):
            try:
                pages.append((i, p.extract_text() or ""))
            except Exception:
                pages.append((i, ""))

        # Extract tables
        pdf_tables = _extract_tables_from_pdf(path)

        # For each page, create a text block; tables will be separate blocks.
        for i, t in pages:
            blocks.append({"type": "text", "page": i, "heading": "", "text": t})

        # Add table blocks
        for idx, t in enumerate(pdf_tables):
            blocks.append({
                "type": "table",
                "page": int(t.get("page", 1)),
                "heading": "",
                "rows": t["rows"],
                "n_cols": int(t.get("n_cols", 0)),
                "table_id": f"table:{idx+1}"
            })

    elif ext in {".docx"}:
        # Body text
        d = Docx(path)
        body_text = "\n".join(p.text or "" for p in d.paragraphs)
        blocks.append({"type": "text", "page": 1, "heading": "", "text": body_text})
        # Tables
        docx_tables = _extract_tables_from_docx(path)
        for idx, t in enumerate(docx_tables):
            blocks.append({
                "type": "table",
                "page": 1,
                "heading": "",
                "rows": t["rows"],
                "n_cols": int(t.get("n_cols", 0)),
                "table_id": f"table:{idx+1}"
            })

    elif ext in {".md", ".markdown"}:
        text = open(path, "r", encoding="utf-8", errors="ignore").read()
        # Detect tables
        md_tables = _detect_md_tables(text)
        # Remove table segments from text for a cleaner text block
        lines = text.splitlines()
        removed = set()
        for t in md_tables:
            s, e = t["span"]
            removed.update(range(s, e))
        text_only = "\n".join(ln for i, ln in enumerate(lines) if i not in removed)
        if text_only.strip():
            blocks.append({"type": "text", "page": 1, "heading": "", "text": text_only})
        for idx, t in enumerate(md_tables):
            blocks.append({
                "type": "table",
                "page": 1,
                "heading": "",
                "rows": t["rows"],
                "n_cols": int(t.get("n_cols", 0)),
                "table_id": f"table:{idx+1}"
            })

    elif ext in {".txt"}:
        text = open(path, "r", encoding="utf-8", errors="ignore").read()
        blocks.append({"type": "text", "page": 1, "heading": "", "text": text})

    else:
        raise ValueError(f"Unsupported format: {ext}")

    return blocks


def tokenize_len(s: str) -> int:
    return max(1, int(len(s.split()) * 1.3))

def split_structure(text: str) -> List[Dict[str, Any]]:
    page = 1
    blocks = []
    current_head = []
    current = []
    lines = text.splitlines()
    for ln in lines:
        if ln.startswith("[[PAGE:"):
            if current:
                blocks.append({"page": page, "heading": " / ".join(current_head[-3:]), "text": "\\n".join(current).strip()})
                current = []
            try:
                page = int(ln.split(":")[1].split("]")[0])
            except Exception:
                pass
            continue
        if ln.strip() and (ln.strip().endswith(":") or (ln.isupper() and len(ln.strip()) < 80)):
            current_head.append(ln.strip().strip(":"))
        current.append(ln)
    if current:
        blocks.append({"page": page, "heading": " / ".join(current_head[-3:]), "text": "\\n".join(current).strip()})
    return blocks

def split_text_structure(text: str) -> List[Dict[str, Any]]:
    page = 1
    blocks = []
    current_head = []
    current = []
    lines = text.splitlines()
    for ln in lines:
        if ln.startswith("[[PAGE:"):
            if current:
                blocks.append({"page": page, "heading": " / ".join(current_head[-3:]), "text": "\n".join(current).strip()})
                current = []
            try:
                page = int(ln.split(":")[1].split("]")[0])
            except Exception:
                pass
            continue
        if ln.strip() and (ln.strip().endswith(":") or (ln.isupper() and len(ln.strip()) < 80)):
            current_head.append(ln.strip().strip(":"))
        current.append(ln)
    if current:
        blocks.append({"page": page, "heading": " / ".join(current_head[-3:]), "text": "\n".join(current).strip()})
    # Tag as type=text
    for b in blocks:
        b["type"] = "text"
    return blocks


def chunk_blocks(mixed_blocks: List[Dict[str, Any]],
                 min_tok=Cfg.target_tokens_min,
                 max_tok=Cfg.target_tokens_max,
                 overlap=Cfg.overlap_ratio) -> List[Dict[str, Any]]:
    """
    Chunk text and table blocks separately.
    Output chunks are normalized to have a 'text' field used for embedding + search,
    and a 'content_type' field in {"text","table"}.
    For tables we also include 'table_id' and 'table_rows' in payload.
    """
    chunks = []

    # 1) Text blocks → break into structural blocks first, then pack into chunks (your original logic)
    text_blocks = [b for b in mixed_blocks if b["type"] == "text" and (b.get("text") or "").strip()]
    # Expand text blocks by structure
    structured = []
    for tb in text_blocks:
        structured.extend(split_text_structure(tb["text"]))

    # Original rolling pack with overlap (adapted)
    buf, buf_pages, buf_headings, cur_tokens = [], [], [], 0

    def flush_text():
        nonlocal buf, buf_pages, buf_headings, cur_tokens
        if not buf:
            return
        text = "\n".join(buf).strip()
        if text:
            chunks.append({
                "content_type": "text",
                "text": text,
                "pages": sorted(set(buf_pages)),
                "heading": " / ".join([h for h in buf_headings if h])[:300]
            })
        buf, buf_pages, buf_headings, cur_tokens = [], [], [], 0

    for b in structured:
        btxt = b["text"].strip()
        if not btxt:
            continue
        btok = tokenize_len(btxt)
        if cur_tokens + btok > max_tok and cur_tokens >= min_tok:
            old = "\n".join(buf)
            keep_chars = int(len(old) * overlap)
            tail = old[-keep_chars:]
            flush_text()
            if tail.strip():
                buf = [tail]
                buf_pages = []
                buf_headings = []
                cur_tokens = tokenize_len(tail)
        buf.append(btxt)
        buf_pages.extend([b.get("page", 1)])
        buf_headings.append(b.get("heading", ""))
        cur_tokens += btok
    flush_text()

    # 2) Table blocks → chunk by rows
    table_blocks = [b for b in mixed_blocks if b["type"] == "table" and b.get("rows")]
    for tb in table_blocks:
        row_chunks = _chunk_table_rows(
            tb["rows"],
            target_tokens_min=max(8, int(min_tok * 0.25)),     # smaller min for tables
            target_tokens_max=max(32, int(max_tok * 0.5))      # cap table chunks smaller
        )
        for i, rch in enumerate(row_chunks):
            md = _render_table_markdown(rch, max_cols=min(tb.get("n_cols", 0) or 0, 12))  # trim extra-wide tables
            if not md.strip():
                continue
            chunks.append({
                "content_type": "table",
                "text": md,                    # for embedding / BM25
                "table_rows": rch,             # stored for payload/debug
                "pages": [tb.get("page", 1)],
                "heading": tb.get("heading", ""),
                "table_id": tb.get("table_id", ""),
                "table_chunk_index": i,
            })

    # Filter trivial chunks and assign chunk_index
    out = [c for c in chunks if tokenize_len(c["text"]) >= 4]
    for i, c in enumerate(out):
        c["chunk_index"] = i
    return out


# ---- Table extraction ----
def _sanitize_cell(x) -> str:
    return (str(x) if x is not None else "").replace("\t", " ").strip()

def _render_table_markdown(rows: List[List[str]], max_cols: int = None) -> str:
    """Render a table to compact GitHub-style Markdown. (max_cols trims wide tables.)"""
    if not rows:
        return ""
    if max_cols is not None:
        rows = [r[:max_cols] for r in rows]
    cols = max(len(r) for r in rows)
    head = rows[0] if rows else []
    # header + separator + rest
    header = "| " + " | ".join(_sanitize_cell(c) or " " for c in head) + " |"
    sep = "| " + " | ".join(["---"] * len(head)) + " |" if head else ""
    body = []
    for r in rows[1:]:
        body.append("| " + " | ".join(_sanitize_cell(c) for c in r) + " |")
    parts = [header, sep] + body if head else body
    return "\n".join([p for p in parts if p])


def _chunk_table_rows(rows: List[List[str]], target_tokens_min: int, target_tokens_max: int) -> List[List[List[str]]]:
    """
    Group table rows into chunks by approximated token count.
    Simple heuristic: token ~ words in the row joined text.
    First row assumed header (keep it in every chunk if present).
    """
    if not rows:
        return []

    header = rows[0]
    data_rows = rows[1:] if len(rows) > 1 else []
    chunks, buf = [], []
    cur_tokens = 0

    def row_tokens(r):  # very cheap proxy
        return max(1, int(len(" ".join(r).split()) * 1.2))

    # Always seed with header (if exists)
    def start_new_chunk():
        nonlocal buf, cur_tokens
        buf = [header[:] ] if header else []
        cur_tokens = row_tokens(header) if header else 0

    start_new_chunk()

    for r in data_rows:
        rt = row_tokens(r)
        if cur_tokens + rt > target_tokens_max and cur_tokens >= target_tokens_min and len(buf) > (1 if header else 0):
            chunks.append(buf)
            start_new_chunk()
        buf.append(r)
        cur_tokens += rt

    if len(buf) > (1 if header else 0):
        chunks.append(buf)
    return chunks

def _extract_tables_from_pdf(path: str):
    """
    Try Camelot first (lattice & stream), then pdfplumber as fallback.
    Returns list of dicts: {"page": int, "rows": List[List[str]], "n_cols": int}
    """
    tables = []
    try:
        # Try both modes; you can conf to use only one for speed
        for flavor in ("lattice", "stream"):
            try:
                tbs = camelot.read_pdf(path, pages="all", flavor=flavor)
                for t in tbs:
                    rows = [[_sanitize_cell(c) for c in row] for row in t.df.values.tolist()]
                    tables.append({"page": t.page, "rows": rows, "n_cols": len(rows[0]) if rows else 0})
            except Exception:
                continue
    except Exception:
        # camelot not available or failed; fall back
        pass

    if not tables:
        try:
            with pdfplumber.open(path) as pdf:
                for i, page in enumerate(pdf.pages, start=1):
                    try:
                        for tbl in page.extract_tables() or []:
                            rows = [[_sanitize_cell(c) for c in row] for row in tbl]
                            tables.append({"page": i, "rows": rows, "n_cols": len(rows[0]) if rows else 0})
                    except Exception:
                        continue
        except Exception:
            pass

    return tables

def _extract_tables_from_docx(path: str):
    """
    Extract DOCX tables via python-docx. Returns same shape as PDF extractor.
    """
    from docx import Document as Docx
    d = Docx(path)
    out = []
    for tbl in d.tables:
        rows = []
        for tr in tbl.rows:
            rows.append([_sanitize_cell(tc.text) for tc in tr.cells])
        out.append({"page": 1, "rows": rows, "n_cols": len(rows[0]) if rows else 0})
    return out

def _detect_md_tables(text: str):
    """
    Very lightweight GFM table detector: captures contiguous pipe-table blocks.
    Returns list of {"page":1, "rows":[...], "n_cols":int, "span":(start,end)}
    """
    import re
    lines = text.splitlines()
    blocks, start = [], None

    def is_table_line(s):
        s = s.strip()
        return s.startswith("|") and s.endswith("|") and "|" in s

    for i, ln in enumerate(lines):
        if is_table_line(ln):
            if start is None:
                start = i
        else:
            if start is not None:
                blocks.append((start, i))
                start = None
    if start is not None:
        blocks.append((start, len(lines)))

    tables = []
    for s, e in blocks:
        seg = lines[s:e]
        # parse rows by | splitting
        rows = []
        for row in seg:
            cells = [c.strip() for c in row.strip().strip("|").split("|")]
            rows.append([_sanitize_cell(c) for c in cells])
        tables.append({"page": 1, "rows": rows, "n_cols": max((len(r) for r in rows), default=0), "span": (s, e)})
    return tables


# Qdrant and Whoosh helpers
def ensure_qdrant():
    client = QdrantClient(
        host=Cfg.qdrant_host,
        port=Cfg.qdrant_port,
        timeout=getattr(Cfg, "qdrant_timeout", 10.0),
    )

    try:
        resp = client.get_collections()
        collections = [c.name for c in resp.collections]
    except Exception as e:
        raise RuntimeError(
            f"Could not connect to Qdrant at {Cfg.qdrant_host}:{Cfg.qdrant_port}. Is it running?"
        ) from e

    if Cfg.collection not in collections:
        client.recreate_collection(
            collection_name=Cfg.collection,
            vectors_config=VectorParams(size=Cfg.dim, distance=Distance.COSINE),
            # Use the *Diff* model so you only set what you care about:
            hnsw_config=HnswConfigDiff(m=Cfg.hnsw_M, ef_construct=Cfg.hnsw_ef_construct),
        )
    return client

def ensure_whoosh(index_dir=Cfg.index_dir):
    os.makedirs(index_dir, exist_ok=True)
    schema = Schema(doc_id=ID(stored=True), chunk_id=ID(stored=True, unique=True),
                    title=TEXT(stored=True), section=TEXT(stored=True),
                    content=TEXT(stored=True), page=TEXT(stored=True))
    if not windex.exists_in(index_dir):
        return windex.create_in(index_dir, schema)
    return windex.open_dir(index_dir)

_embedder = None
def embedder():
    global _embedder
    if _embedder is None:
        _embedder = SentenceTransformer(Cfg.embed_model, device="cpu")
    return _embedder

def upsert(qc: QdrantClient, wi, doc_path: str, doc_title: str = None, doc_id: str = None):
    mixed_blocks = read_blocks_with_tables(doc_path)
    doc_id = doc_id or sha1(doc_path)
    doc_title = doc_title or os.path.basename(doc_path)

    chunks = chunk_blocks(mixed_blocks)
    if not chunks:
        return {"doc_id": doc_id, "chunks": 0, "reason": "no_chunks"}

    model = embedder()
    vectors = model.encode([c["text"] for c in chunks], normalize_embeddings=True, show_progress_bar=True)
    if len(vectors) == 0:
        return {"doc_id": doc_id, "chunks": 0, "reason": "no_vectors"}

    dim = len(vectors[0])
    if dim != Cfg.dim:
        raise ValueError(f"Embedding dim {dim} != configured Cfg.dim={Cfg.dim}. Set Cfg.dim to {dim} or switch model.")

    ids = [str(uuid.uuid5(uuid.NAMESPACE_URL, f"{doc_id}:{c['chunk_index']}")) for c in chunks]

    payloads = []
    for c in chunks:
        payloads.append({
            "doc_id": doc_id,
            "doc_title": doc_title,
            "source_path": os.path.abspath(doc_path),
            "section_title": c.get("heading", ""),
            "page_nums": c.get("pages", []),
            "chunk_index": c["chunk_index"],
            "content_type": c.get("content_type", "text"),        # NEW
            "table_id": c.get("table_id", ""),                    # NEW
            "table_chunk_index": c.get("table_chunk_index", -1),  # NEW
            "n_table_rows": len(c.get("table_rows", [])),         # NEW
            "created_at": now(),
        })

    batch = models.Batch.construct(
        ids=ids,
        vectors=[v.tolist() for v in vectors],
        payloads=payloads
    )
    qc.upsert(collection_name=Cfg.collection, points=batch, wait=True)

    # Whoosh: store a simple textual representation regardless of type (so BM25 works)
    writer = wi.writer()
    for c in chunks:
        cid = f"{doc_id}:{c['chunk_index']}"
        writer.update_document(
            doc_id=doc_id,
            chunk_id=cid,
            title=doc_title,
            section=c.get("heading", ""),
            content=c["text"],                  # tables are markdown text → BM25-able
            page=",".join(map(str, c.get("pages", [])))
        )
    writer.commit()
    return {"doc_id": doc_id, "chunks": len(chunks), "reason": "ok"}
